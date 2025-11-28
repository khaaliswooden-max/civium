"""
Document Extraction Module
==========================

Extracts text content from various document formats:
- PDF documents
- HTML web pages
- Plain text files
- Word documents (docx)

Version: 0.1.0
"""

import hashlib
import io
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from enum import Enum
from pathlib import Path
from typing import Any

import httpx

from shared.logging import get_logger

logger = get_logger(__name__)


class DocumentFormat(str, Enum):
    """Supported document formats."""

    PDF = "pdf"
    HTML = "html"
    TEXT = "text"
    MARKDOWN = "markdown"
    DOCX = "docx"
    UNKNOWN = "unknown"


@dataclass
class ExtractionResult:
    """Result of document extraction."""

    content: str
    format: DocumentFormat
    source_url: str | None = None
    source_path: str | None = None

    # Metadata
    title: str | None = None
    page_count: int | None = None
    word_count: int = 0
    char_count: int = 0

    # Hashing for change detection
    content_hash: str = ""

    # Timestamps
    extracted_at: datetime = field(default_factory=lambda: datetime.now(UTC))

    # Additional metadata
    metadata: dict[str, Any] = field(default_factory=dict)

    # Extraction quality indicators
    confidence: float = 1.0
    warnings: list[str] = field(default_factory=list)

    def __post_init__(self) -> None:
        """Calculate derived fields."""
        if not self.content_hash and self.content:
            self.content_hash = hashlib.sha256(self.content.encode()).hexdigest()
        if not self.word_count and self.content:
            self.word_count = len(self.content.split())
        if not self.char_count and self.content:
            self.char_count = len(self.content)


class DocumentExtractor:
    """
    Extracts text content from various document formats.

    Supports:
    - PDF extraction with layout preservation
    - HTML cleaning and text extraction
    - Plain text normalization
    - URL fetching with retry logic
    """

    def __init__(
        self,
        timeout: float = 30.0,
        max_retries: int = 3,
        user_agent: str = "CiviumBot/1.0 (Regulatory Intelligence)",
    ) -> None:
        """
        Initialize the document extractor.

        Args:
            timeout: HTTP request timeout in seconds
            max_retries: Maximum retry attempts for failed requests
            user_agent: User agent string for HTTP requests
        """
        self.timeout = timeout
        self.max_retries = max_retries
        self.user_agent = user_agent

        self._http_client: httpx.AsyncClient | None = None

    async def _get_client(self) -> httpx.AsyncClient:
        """Get or create HTTP client."""
        if self._http_client is None:
            self._http_client = httpx.AsyncClient(
                timeout=httpx.Timeout(self.timeout),
                headers={"User-Agent": self.user_agent},
                follow_redirects=True,
            )
        return self._http_client

    async def close(self) -> None:
        """Close the HTTP client."""
        if self._http_client:
            await self._http_client.aclose()
            self._http_client = None

    def detect_format(self, content: bytes | str, url: str | None = None) -> DocumentFormat:
        """
        Detect document format from content or URL.

        Args:
            content: Document content (bytes or string)
            url: Source URL (optional)

        Returns:
            Detected DocumentFormat
        """
        # Check URL extension first
        if url:
            url_lower = url.lower()
            if url_lower.endswith(".pdf"):
                return DocumentFormat.PDF
            elif url_lower.endswith((".html", ".htm")):
                return DocumentFormat.HTML
            elif url_lower.endswith(".md"):
                return DocumentFormat.MARKDOWN
            elif url_lower.endswith(".docx"):
                return DocumentFormat.DOCX
            elif url_lower.endswith(".txt"):
                return DocumentFormat.TEXT

        # Check content magic bytes for binary formats
        if isinstance(content, bytes):
            if content.startswith(b"%PDF"):
                return DocumentFormat.PDF
            elif content.startswith(b"PK"):  # ZIP-based (docx, xlsx, etc.)
                return DocumentFormat.DOCX

        # Check content for HTML markers
        content_str = content if isinstance(content, str) else content.decode("utf-8", errors="ignore")
        if "<html" in content_str.lower()[:1000] or "<!doctype html" in content_str.lower()[:1000]:
            return DocumentFormat.HTML

        return DocumentFormat.TEXT

    async def extract_from_url(self, url: str) -> ExtractionResult:
        """
        Extract document content from a URL.

        Args:
            url: URL to fetch and extract

        Returns:
            ExtractionResult with extracted content
        """
        client = await self._get_client()
        warnings: list[str] = []

        for attempt in range(self.max_retries):
            try:
                response = await client.get(url)
                response.raise_for_status()

                content_type = response.headers.get("content-type", "").lower()

                # Determine format from content type
                if "pdf" in content_type:
                    doc_format = DocumentFormat.PDF
                elif "html" in content_type:
                    doc_format = DocumentFormat.HTML
                else:
                    doc_format = self.detect_format(response.content, url)

                # Extract based on format
                if doc_format == DocumentFormat.PDF:
                    result = await self._extract_pdf(response.content)
                elif doc_format == DocumentFormat.HTML:
                    result = self._extract_html(response.text)
                else:
                    result = self._extract_text(response.text)

                result.source_url = url
                result.format = doc_format
                result.warnings.extend(warnings)

                logger.info(
                    "document_extracted_from_url",
                    url=url,
                    format=doc_format.value,
                    chars=result.char_count,
                )

                return result

            except httpx.HTTPStatusError as e:
                logger.warning(
                    "http_error",
                    url=url,
                    status=e.response.status_code,
                    attempt=attempt + 1,
                )
                if attempt == self.max_retries - 1:
                    raise
            except Exception as e:
                logger.error(
                    "extraction_error",
                    url=url,
                    error=str(e),
                    attempt=attempt + 1,
                )
                if attempt == self.max_retries - 1:
                    raise

        raise RuntimeError(f"Failed to extract from {url} after {self.max_retries} attempts")

    async def extract_from_file(self, file_path: str | Path) -> ExtractionResult:
        """
        Extract document content from a file.

        Args:
            file_path: Path to the file

        Returns:
            ExtractionResult with extracted content
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        content = path.read_bytes()
        doc_format = self.detect_format(content, str(path))

        if doc_format == DocumentFormat.PDF:
            result = await self._extract_pdf(content)
        elif doc_format == DocumentFormat.HTML:
            result = self._extract_html(content.decode("utf-8", errors="replace"))
        elif doc_format == DocumentFormat.DOCX:
            result = self._extract_docx(content)
        else:
            result = self._extract_text(content.decode("utf-8", errors="replace"))

        result.source_path = str(path)
        result.format = doc_format

        logger.info(
            "document_extracted_from_file",
            path=str(path),
            format=doc_format.value,
            chars=result.char_count,
        )

        return result

    async def extract_from_bytes(
        self,
        content: bytes,
        filename: str | None = None,
    ) -> ExtractionResult:
        """
        Extract document content from raw bytes.

        Args:
            content: Document content as bytes
            filename: Optional filename for format detection

        Returns:
            ExtractionResult with extracted content
        """
        doc_format = self.detect_format(content, filename)

        if doc_format == DocumentFormat.PDF:
            return await self._extract_pdf(content)
        elif doc_format == DocumentFormat.HTML:
            return self._extract_html(content.decode("utf-8", errors="replace"))
        elif doc_format == DocumentFormat.DOCX:
            return self._extract_docx(content)
        else:
            return self._extract_text(content.decode("utf-8", errors="replace"))

    async def _extract_pdf(self, content: bytes) -> ExtractionResult:
        """
        Extract text from PDF content.

        Uses pypdf for extraction with fallback strategies.
        """
        warnings: list[str] = []
        text_parts: list[str] = []
        page_count = 0
        title: str | None = None

        try:
            # Try pypdf first
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)

            # Extract metadata
            if reader.metadata:
                title = reader.metadata.get("/Title")

            # Extract text from each page
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
                except Exception as e:
                    warnings.append(f"Failed to extract page {i + 1}: {str(e)}")

            text = "\n\n".join(text_parts)

            # If no text extracted, might be scanned PDF
            if not text.strip():
                warnings.append("No text extracted - document may be scanned/image-based")

        except ImportError:
            warnings.append("pypdf not installed, using fallback extraction")
            text = "[PDF extraction requires pypdf package]"
        except Exception as e:
            logger.error("pdf_extraction_error", error=str(e))
            warnings.append(f"PDF extraction error: {str(e)}")
            text = ""

        return ExtractionResult(
            content=text,
            format=DocumentFormat.PDF,
            title=title,
            page_count=page_count,
            warnings=warnings,
            confidence=0.9 if text else 0.1,
        )

    def _extract_html(self, content: str) -> ExtractionResult:
        """
        Extract text from HTML content.

        Uses BeautifulSoup for parsing with text cleaning.
        """
        warnings: list[str] = []
        title: str | None = None

        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(content, "lxml")

            # Extract title
            title_tag = soup.find("title")
            if title_tag:
                title = title_tag.get_text(strip=True)

            # Remove script and style elements
            for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
                element.decompose()

            # Find main content area if present
            main_content = soup.find("main") or soup.find("article") or soup.find(
                "div", {"class": re.compile(r"content|main|body", re.I)}
            )

            if main_content:
                text = main_content.get_text(separator="\n", strip=True)
            else:
                text = soup.get_text(separator="\n", strip=True)

            # Clean up whitespace
            text = re.sub(r"\n{3,}", "\n\n", text)
            text = re.sub(r" {2,}", " ", text)

        except ImportError:
            warnings.append("beautifulsoup4/lxml not installed, using regex fallback")
            # Fallback: basic regex extraction
            text = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL | re.I)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.I)
            text = re.sub(r"<[^>]+>", " ", text)
            text = re.sub(r"\s+", " ", text).strip()

        except Exception as e:
            logger.error("html_extraction_error", error=str(e))
            warnings.append(f"HTML extraction error: {str(e)}")
            text = content

        return ExtractionResult(
            content=text,
            format=DocumentFormat.HTML,
            title=title,
            warnings=warnings,
        )

    def _extract_text(self, content: str) -> ExtractionResult:
        """
        Process plain text content.

        Normalizes whitespace and line endings.
        """
        # Normalize line endings
        text = content.replace("\r\n", "\n").replace("\r", "\n")

        # Normalize whitespace
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Detect and extract title from first line if it looks like a title
        lines = text.strip().split("\n")
        title: str | None = None
        if lines and len(lines[0]) < 200 and not lines[0].endswith((".", ",", ";")):
            title = lines[0].strip()

        return ExtractionResult(
            content=text.strip(),
            format=DocumentFormat.TEXT,
            title=title,
        )

    def _extract_docx(self, content: bytes) -> ExtractionResult:
        """
        Extract text from DOCX content.

        Uses python-docx for extraction.
        """
        warnings: list[str] = []
        title: str | None = None
        text_parts: list[str] = []

        try:
            from docx import Document

            doc = Document(io.BytesIO(content))

            # Extract core properties for title
            if doc.core_properties.title:
                title = doc.core_properties.title

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            text = "\n\n".join(text_parts)

        except ImportError:
            warnings.append("python-docx not installed")
            text = "[DOCX extraction requires python-docx package]"
        except Exception as e:
            logger.error("docx_extraction_error", error=str(e))
            warnings.append(f"DOCX extraction error: {str(e)}")
            text = ""

        return ExtractionResult(
            content=text,
            format=DocumentFormat.DOCX,
            title=title,
            warnings=warnings,
        )

